import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import re

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run element
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add color and underline
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    r.append(rPr)
    r_text = OxmlElement('w:t')
    r_text.text = text
    r.append(r_text)
    hyperlink.append(r)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def process_line_with_links(paragraph, line):
    """Process a line that may contain markdown links and bold text."""
    # Pattern to find markdown links
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    # Pattern to find bold text
    bold_pattern = r'\*\*([^\*]+)\*\*'
    
    # Split the line by links and bold text
    parts = []
    last_end = 0
    
    # Find all links and bold text
    all_matches = []
    for match in re.finditer(link_pattern, line):
        all_matches.append(('link', match))
    for match in re.finditer(bold_pattern, line):
        all_matches.append(('bold', match))
    
    # Sort matches by position
    all_matches.sort(key=lambda x: x[1].start())
    
    for match_type, match in all_matches:
        # Add text before the match
        if match.start() > last_end:
            text_before = line[last_end:match.start()]
            if text_before:
                paragraph.add_run(text_before)
        
        if match_type == 'link':
            text = match.group(1)
            url = match.group(2)
            add_hyperlink(paragraph, text, url)
        elif match_type == 'bold':
            text = match.group(1)
            paragraph.add_run(text).bold = True
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(line):
        paragraph.add_run(line[last_end:])

# Read the JSON file
with open('enquiries@singita.com_2025_06_guest-post-where-the-light-begins_ copy.json', 'r') as f:
    data = json.load(f)

# Create a new Document
doc = Document()

# Add the URL at the top
url = data['metadata']['ogUrl']
url_paragraph = doc.add_paragraph()
url_paragraph.add_run(url).bold = True
url_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
doc.add_paragraph()  # Add empty line after URL

# Get the full markdown content
markdown_content = data['markdown']

# Process the content line by line
lines = markdown_content.split('\n')

for line in lines:
    if not line.strip():
        doc.add_paragraph()
        continue
    
    # Handle headers
    if line.startswith('# '):
        p = doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=4)
    # Handle italic text
    elif line.strip().startswith('_') and line.strip().endswith('_'):
        p = doc.add_paragraph()
        p.add_run(line.strip()[1:-1]).italic = True
    # Handle image links
    elif '![' in line and '](' in line:
        p = doc.add_paragraph()
        # Extract image description and URL
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            img_desc = img_match.group(1)
            img_url = img_match.group(2)
            if img_desc:
                p.add_run(f"[Image: {img_desc}]").italic = True
            else:
                p.add_run("[Image]").italic = True
            p.add_run(f"\n{img_url}")
    # Handle lines with links or bold text
    elif ('[' in line and '](' in line and not line.strip().startswith('!')) or '**' in line:
        p = doc.add_paragraph()
        process_line_with_links(p, line)
    # Handle regular paragraphs
    else:
        p = doc.add_paragraph()
        p.add_run(line)

# Save the document
doc.save('guest-post-light-begins.docx')
print("Document created successfully: guest-post-light-begins.docx")